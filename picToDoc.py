import os
from docx import Document
from docx.shared import Inches
from PIL import Image, ExifTags

# Konfiguration
BILDER_ORDNER = "Pics"
OPTIMIERTER_ORDNER = "optimierte_bilder"
OUTPUT_DATEI = "Fotoprotokoll.docx"
MAX_BREITE_INCHES = 6  # Maximale Bildbreite im Dokument (angepasst)
DPI = 300  # Standard DPI für Word, um echte Zoll-Werte zu erhalten
bildCounter = 0

# Stelle sicher, dass der Ordner für optimierte Bilder existiert
if not os.path.exists(OPTIMIERTER_ORDNER):
    os.makedirs(OPTIMIERTER_ORDNER)

# Funktion: Exif-Daten entfernen, drehen & Bildgröße anpassen
def verarbeite_bild(bildpfad):
    with Image.open(bildpfad) as img:
        # 1. Exif-Daten prüfen und Bild richtig drehen
        try:
            for tag in ExifTags.TAGS:
                if ExifTags.TAGS[tag] == "Orientation":
                    orientation_tag = tag
                    break
            exif = img._getexif()
            if exif and orientation_tag in exif:
                orientation = exif[orientation_tag]
                if orientation == 3:
                    img = img.rotate(180, expand=True)
                elif orientation == 6:
                    img = img.rotate(270, expand=True)
                elif orientation == 8:
                    img = img.rotate(90, expand=True)
        except (AttributeError, KeyError, IndexError):
            pass  # Falls Exif-Daten fehlen, nichts tun

        # 2. Exif-Daten entfernen, damit Word keine falschen Drehungen übernimmt
        img = img.convert("RGB")  # Konvertiere das Bild, um Exif zu entfernen
        img.info["exif"] = None  # Lösche Exif-Daten

        # 3. Größe anpassen, aber nicht zu stark verkleinern
        breite, hoehe = img.size
        max_pixel_breite = int(MAX_BREITE_INCHES * DPI)  # 5 Zoll x 96 DPI
        faktor = max_pixel_breite / breite
        if faktor < 1:  # Nur verkleinern, nicht vergrößern
            neue_breite = int(breite * faktor)
            neue_hoehe = int(hoehe * faktor)
            img = img.resize((neue_breite, neue_hoehe), Image.LANCZOS)

        # 4. Bearbeitetes Bild im separaten Ordner speichern
        bildname = os.path.basename(bildpfad)
        optimiertes_bildpfad = os.path.join(OPTIMIERTER_ORDNER, bildname)
        img.save(optimiertes_bildpfad, "JPEG", quality=95)  # Qualität beibehalten
        return optimiertes_bildpfad


# 1. Erstelle ein Word-Dokument
doc = Document()
doc.add_heading("Fotoprotokoll", level=1)

# 2. Durchlaufe alle Bilder im Ordner
if not os.path.exists(BILDER_ORDNER):
    print(f"Ordner '{BILDER_ORDNER}' nicht gefunden!")
    exit()

bilder = sorted(os.listdir(BILDER_ORDNER))  # Bilder sortieren


for bild in bilder:
    if bild.lower().endswith((".png", ".jpg", ".jpeg")):
        bildpfad = os.path.join(BILDER_ORDNER, bild)
        optimiertes_bild = verarbeite_bild(bildpfad)  # Korrigiertes Bild holen

        doc.add_paragraph(f"Bild: {bild}")
        
        # **Sicherstellen, dass das optimierte Bild existiert**
        if os.path.exists(optimiertes_bild):
            doc.add_picture(optimiertes_bild, width=Inches(MAX_BREITE_INCHES))
        else:
            print(f"Warnung: Verkleinertes Bild {optimiertes_bild} nicht gefunden, Original wird genutzt!")
            doc.add_picture(bildpfad, width=Inches(MAX_BREITE_INCHES))

        doc.add_paragraph("")  # Leerzeile für bessere Lesbarkeit
        bildCounter += 1

doc.add_paragraph(f"Anzahl Bilder: " + str(bildCounter))

# 3. Speichere das Word-Dokument
doc.save(OUTPUT_DATEI)
print(f"Dokument gespeichert als {OUTPUT_DATEI}")
print(f"Optimierte Bilder wurden gespeichert in '{OPTIMIERTER_ORDNER}'")

# 4. (Optional) Word-Datei in PDF umwandeln
try:
    import comtypes.client
    wdFormatPDF = 17
    word = comtypes.client.CreateObject("Word.Application")
    doc = word.Documents.Open(os.path.abspath(OUTPUT_DATEI))
    doc.SaveAs(os.path.abspath("Fotoprotokoll.pdf"), FileFormat=wdFormatPDF)
    doc.Close()
    word.Quit()
    print("PDF erfolgreich erstellt: Fotoprotokoll.pdf")
except ImportError:
    print("PDF-Export nicht möglich. Installiere 'comtypes', falls du Word nutzt.")
